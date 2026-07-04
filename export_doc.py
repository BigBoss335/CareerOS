#!/usr/bin/env python3

import argparse
import re
from docx import Document
from docx.shared import Pt


def add_bold_paragraph(paragraph, text):
    """
    Convert **bold** markdown to Word bold runs.
    """
    parts = re.split(r'(\*\*.*?\*\*)', text)

    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def markdown_to_docx(input_file, output_file):
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    with open(input_file, "r", encoding="utf-8") as f:
        lines = f.readlines()

    for line in lines:

        line = line.rstrip()

        if not line:
            doc.add_paragraph()
            continue

        # H1
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
            continue

        # H2
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
            continue

        # Bullet list
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_bold_paragraph(p, line[2:])
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        add_bold_paragraph(p, line)

    doc.save(output_file)


def main():
    parser = argparse.ArgumentParser(
        description="Convert CareerOS Markdown to DOCX"
    )

    parser.add_argument(
        "--input",
        required=True,
        help="Input Markdown file"
    )

    parser.add_argument(
        "--output",
        required=True,
        help="Output DOCX file"
    )

    args = parser.parse_args()

    markdown_to_docx(args.input, args.output)

    print(f"Created: {args.output}")

if __name__ == "__main__":
    main()
